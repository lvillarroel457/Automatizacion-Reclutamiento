from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from main.domain.cv import CV
from main.ports.cv_formatter import CVFormatter

class DOCXCVFormatter(CVFormatter):

    def format(self, cv: CV, output_file: str):
        doc = Document()
        
        # Add candidate name
        p = doc.add_paragraph()
        run = p.add_run(cv.name)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Proxima Nova'
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add candidate position
        p = doc.add_paragraph()
        run = p.add_run(cv.position)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Open Sans'
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add Summary
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells

        # Format "Summary of Qualifications" cell
        hdr_run_0 = hdr_cells[0].paragraphs[0].add_run('Summary of Qualifications')
        hdr_run_0.font.size = Pt(12)
        hdr_run_0.font.bold = True
        hdr_run_0.font.name = 'Open Sans'
        hdr_cells[0].width = Pt(150)
        hdr_cells[0].text_alignment = WD_ALIGN_PARAGRAPH.LEFT
        hdr_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Format the summary text cell
        hdr_run_1 = hdr_cells[1].paragraphs[0].add_run(cv.summary)
        hdr_run_1.font.size = Pt(9)
        hdr_run_1.font.name = 'Open Sans'
        hdr_cells[1].text_alignment = WD_ALIGN_PARAGRAPH.LEFT
        hdr_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Skills
        skills_para = doc.add_paragraph()
        run = skills_para.add_run('Skills')
        run.bold = True
        skills_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for skill_category in cv.skills:
            cat_para = doc.add_paragraph()
            run = cat_para.add_run(skill_category.name)
            run.bold = True
            for skill in skill_category.skills:
                doc.add_paragraph(skill)
                
        # Experience
        experience_para = doc.add_paragraph()
        run = experience_para.add_run('Experience')
        run.bold = True
        experience_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for project in cv.experience:
            doc.add_paragraph().add_run('Project Name:').bold = True
            doc.add_paragraph(project.name)
            doc.add_paragraph().add_run('Project Description:').bold = True
            doc.add_paragraph(project.description)
            doc.add_paragraph().add_run('Customer:').bold = True
            doc.add_paragraph(project.customer)
            doc.add_paragraph().add_run('Duration:').bold = True
            doc.add_paragraph(project.duration)
            doc.add_paragraph().add_run('Role:').bold = True
            doc.add_paragraph(project.role)
            doc.add_paragraph().add_run('Responsibilities:').bold = True
            doc.add_paragraph(project.responsibilities)
            doc.add_paragraph().add_run('Team Size:').bold = True
            doc.add_paragraph(str(project.team_size))
            doc.add_paragraph().add_run('Tools & Technologies:').bold = True
            doc.add_paragraph(', '.join(project.tools_technologies))
            
        # Save to the specified output file
        doc.save(output_file)

    def get_extension(self):
        return '.docx'
