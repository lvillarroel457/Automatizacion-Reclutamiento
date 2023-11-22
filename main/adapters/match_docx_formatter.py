from main.ports.matcher import MatchFormatter
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate

class DOCXMatchFormatter(MatchFormatter):

    def format_match(self, MatchList, output_file: str):

        ''' Crea un documento .docx con los resultados del match. '''

        doc = Document()

        n=len(MatchList)

# Iterate through the dictionary and add each key-value pair as a line in the paragraph
        for i,c in enumerate(MatchList):

            if i==0:
                p = doc.add_paragraph()
                line = f"Informe Match para el cargo {c['Data']['Cargo']}"
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(20)
                run.font.name = 'Open Sans'
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER              

            for key, value in c.items():
                
                if key=="Data":

                    p = doc.add_paragraph()
                    line = f"Nombre : {value['Nombre']}"
                    run = p.add_run(line)
                    run.bold = True
                    run.font.size = Pt(18)
                    run.font.name = 'Open Sans'
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    p = doc.add_paragraph()
                    line = f" Posición Match: {i+1} de {n}"
                    run = p.add_run(line)
                    run.bold = True
                    run.font.size = Pt(14)
                    run.font.name = 'Open Sans'
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


                    
                    mean = c['Porcentaje Promedio']
                    p = doc.add_paragraph()
                    line = f" El porcentaje promedio es : {mean}"
                    run = p.add_run(line) 
                    run.bold = True
                    run.font.size = Pt(12)
                    run.font.name = 'Open Sans'
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                elif key!="Porcentaje Promedio":
                    p = doc.add_paragraph()
                    line = f"{key}: {value}"
                    run = p.add_run(line)
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Open Sans'
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc.save(output_file)


    def format_matchtpl(self, MatchList, output_file: str):

        ''' Crea un documento .docx con los resultados del match según un template con la librería docxtpl. '''

        n=len(MatchList)
        doc = DocxTemplate('Match_Format.docx')
        context = {'cargo' : MatchList[0]['Data']['Cargo'], 'candidatos' : MatchList, 'total': n }
        doc.render(context)
        doc.save(output_file)



